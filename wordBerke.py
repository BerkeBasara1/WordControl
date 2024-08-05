import docx
import re
import difflib
from collections import Counter
from docx import Document
from docx.shared import RGBColor
class WordTextFormatter:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.text = ""

    def extract_text_from_docx(self):
        doc = docx.Document(self.docx_path)
        for para in doc.paragraphs:
            self.text += para.text + "\n"

    def format_text(self):
        # Temizleme işlemleri
        self.text = re.sub(r'(\w)(\s{2,})(\w)', r'\1 \3', self.text)  # Kelimeler arasında fazla boşlukları düzeltme
        self.text = re.sub(r'\s+', ' ', self.text)  # Fazla boşlukları temizleme
        self.text = re.sub(r'\s([.,;:!?])', r'\1', self.text)  # Noktalama işaretlerinden önceki boşlukları temizleme
        self.text = re.sub(r'([.,;:!?])([A-Za-z])', r'\1 \2', self.text)  # Noktalama işaretlerinden sonraki boşlukları ekleme
        self.text = re.sub(r'(\d)(\s+)(\d)', r'\1\3', self.text)  # Sayılar arasında fazla boşlukları düzeltme
        self.text = re.sub(r'\s+\n', '\n', self.text).strip()  # Satır başındaki boşlukları temizleme
        
        # Paragrafları ayırma
        formatted_text = ""
        lines = self.text.split('\n')
        for line in lines:
            if line.strip():
                formatted_text += line.strip() + ' '
            else:
                formatted_text = formatted_text.strip() + '\n\n'
        
        self.text = re.sub(r'\n\n+', '\n\n', formatted_text)  # Fazla boş satırları temizleme

    def get_formatted_text(self):
        self.extract_text_from_docx()
        self.format_text()
        return self.text.strip()
def highlight_differences(text1, text2):
    diff = difflib.ndiff(text1.split(), text2.split())
    highlighted_text1 = []
    highlighted_text2 = []

    for word in diff:
        if word.startswith('- '):
            highlighted_text1.append(('red', word[2:]))
        elif word.startswith('+ '):
            highlighted_text2.append(('green', word[2:]))
        else:
            word = word[2:]
            highlighted_text1.append(('black', word))
            highlighted_text2.append(('black', word))
    
    return highlighted_text1, highlighted_text2
def compare_docs(docx_path1, docx_path2):
    formatter1 = WordTextFormatter(docx_path1)
    text1 = formatter1.get_formatted_text()
    
    formatter2 = WordTextFormatter(docx_path2)
    text2 = formatter2.get_formatted_text()

    # Benzerlik yüzdesini hesaplama
    sequence_matcher = difflib.SequenceMatcher(None, text1, text2)
    similarity_percentage = sequence_matcher.ratio() * 100

    # Farklı kelimeleri bulma
    words1 = Counter(re.findall(r'\w+', text1))
    words2 = Counter(re.findall(r'\w+', text2))
    
    diff_words1 = words1 - words2
    diff_words2 = words2 - words1
    
    different_words_count = sum(diff_words1.values()) + sum(diff_words2.values())

    # Farklı kelimeleri metinlerde belirterek gösterme
    highlighted_text1, highlighted_text2 = highlight_differences(text1, text2)

    return similarity_percentage, different_words_count, highlighted_text1, highlighted_text2
def create_word(highlighted_text, output_path):
    doc = Document()
    para = doc.add_paragraph()
    
    for color, word in highlighted_text:
        run = para.add_run(word + ' ')
        if color == 'red':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif color == 'green':
            run.font.color.rgb = RGBColor(0, 255, 0)
        # Varsayılan renk siyah (black) olduğundan, başka bir işlem gerekmez.

    doc.save(output_path)
# Word dosyalarının yolları
docx_path1 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge1.docx"
docx_path2 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge2.docx"

# Word dosyalarını kıyaslama
similarity_percentage, different_words_count, highlighted_text1, highlighted_text2 = compare_docs(docx_path1, docx_path2)

# Benzerlik yüzdesini ve farklı kelime sayısını yazdırma
print(f"Benzerlik Yüzdesi: {similarity_percentage:.2f}%")
print(f"Farklı Kelime Sayısı: {different_words_count}")

# Terminalde farkları yazdırma
def print_highlighted_text(highlighted_text):
    for color, word in highlighted_text:
        if color == 'red':
            print(f"\033[91m{word}\033[0m", end=' ')
        elif color == 'green':
            print(f"\033[92m{word}\033[0m", end=' ')
        else:
            print(word, end=' ')
    print()

print("Birinci Word dosyasının farklı kelimeleri:")
print_highlighted_text(highlighted_text1)

print("İkinci Word dosyasının farklı kelimeleri:")
print_highlighted_text(highlighted_text2)

# Yeni Word dosyalarını oluşturma
output_path1 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge3.docx"
output_path2 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge4.docx"
create_word(highlighted_text1, output_path1)
create_word(highlighted_text2, output_path2)
